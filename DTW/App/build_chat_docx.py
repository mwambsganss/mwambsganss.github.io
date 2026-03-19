#!/usr/bin/env python3
"""Standalone helper: build a chat-log .docx from session JSON file.

Usage:
    python3 build_chat_docx.py <session_json_path> '<exercises_json>'
Output: raw .docx bytes written to stdout.
"""
import datetime
import io
import json
import sys

from docx import Document
from docx.shared import Pt, RGBColor

with open(sys.argv[1], encoding="utf-8") as _f:
    session = json.load(_f)
all_ex_meta = json.loads(sys.argv[2])

doc = Document()

title = doc.add_heading(session["session_name"], 0)
title.alignment = 1
sub = doc.add_paragraph(f"Chat Log Export — {datetime.date.today().isoformat()}")
sub.alignment = 1
doc.add_paragraph()

participants = session.get("participants", [])
if participants:
    doc.add_heading("Participants", 2)
    for p in participants:
        line = p.get("name", "Unknown")
        if p.get("role"):
            line += f"  —  {p['role']}"
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph()

ex_order       = session.get("exercise_order", [])
exercises_data = session.get("exercises", {})

for ex_id in ex_order:
    ex_state = exercises_data.get(ex_id, {})
    messages = ex_state.get("messages", [])
    if not messages:
        continue

    ex_meta = next((e for e in all_ex_meta if e["id"] == ex_id), None)
    ex_name = ex_meta["name"] if ex_meta else ex_id.replace("_", " ").title()

    doc.add_heading(ex_name, 1)

    for msg in messages:
        sender = msg.get("sender_name", "Unknown")
        role   = msg.get("role", "")
        text   = msg.get("text", "")
        ts     = msg.get("timestamp", "")

        p = doc.add_paragraph()
        run_name = p.add_run(sender)
        run_name.bold = True
        if role:
            run_name = p.add_run(f"  ({role})")
            run_name.font.size = Pt(9)
            run_name.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        if ts:
            try:
                dt = datetime.datetime.fromisoformat(ts)
                ts_str = dt.strftime("  %I:%M %p")
            except Exception:
                ts_str = f"  {ts}"
            run_ts = p.add_run(ts_str)
            run_ts.font.size = Pt(9)
            run_ts.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        body = doc.add_paragraph(text)
        body.paragraph_format.left_indent = Pt(12)
        body.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

buf = io.BytesIO()
doc.save(buf)
sys.stdout.buffer.write(buf.getvalue())
